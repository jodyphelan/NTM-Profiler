import sys
from time import time
from docxtpl import DocxTemplate
from collections import defaultdict
from .models import ProfileResult, SpeciesResult, VcfQC, FastaQC, BamQC
from docx import Document
from typing import List
from copy import deepcopy
import logging
from abc import abstractmethod
from .utils import rv2genes, get_gene2drugs
import os



class DocxResultTemplate:
    @abstractmethod
    def write_output(self, result: ProfileResult, conf: dict, outfile: str):
        """Write the output to the Word document"""
        pass

def sanitize(d):
    d = d.replace("-","_")
    return d

def cache_cells(tab) -> List[List[str]]:
    _cells = tab._cells
    numcols = len(tab.columns)
    cells = []
    for row in tab.rows:
        rowcells = []
        for i in range(numcols):
            rowcells.append(_cells.pop(0))
        cells.append(rowcells)
    return cells

def merge_cells(filename: str) -> None:
    doc = Document(filename)
    
    def _merge_cells(tab, rows: List[int], column: int):
        
        if column >= len(tab.columns) - 1:
            return
        if len(rows)==1:
            return 
        c1 = tab.cell(rows[0], column)
        c2 = tab.cell(rows[-1], column)
        c1_font_size = c1.paragraphs[0].runs[0].font.size

        for r in rows[1:]:
            tab.cell(r, column).text = ''
     

        cm = c1.merge(c2)


            
        values_in_next_column = set([tab.rows[r].cells[column+1].text for r in rows])
        for val in values_in_next_column:
            rows_with_val = [r for r in rows if tab.c[r][column+1].text == val]
            _merge_cells(tab, rows_with_val, column+1)
    
    for tab in doc.tables:
        tab.c = cache_cells(tab)
        values_in_next_column = set([tab.rows[r].cells[0].text for r in range(1, len(tab.rows))])
        for val in values_in_next_column:
            rows_with_val = [r for r in range(1, len(tab.rows)) if tab.c[r][0].text == val]

            _merge_cells(tab, rows_with_val, 0)

    doc.save(filename)
class DefaultSpeciesTemplate(DocxResultTemplate):
    __template_name__ = "default-species"
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        __dir__ = os.path.dirname(os.path.abspath(__file__))
        self.template_filename = os.path.join(__dir__, "report-templates", "default_species_template.docx")
    def write_output(
            self, 
            result: SpeciesResult, 
            conf: dict,
            output_filename: str
        ):
        species_rows = []
        for t in result.taxa:
            species_rows.append({
                'species': t.species,
                'accession': t.accession,
                'ani': round(t.ani,2) if (t.ani and result.data_source in ('fastq','bam')) else 'N/A',
                'rabund': round(t.relative_abundance,2) if (t.relative_abundance and result.data_source in ('fastq','bam')) else 'N/A',
                'notes': '\n'.join(t.notes)
            })
        if result.data_source in ('fastq','bam'):
            # result.taxa.abundance
            median_depth = ",".join([f"{t.species}: {t.abundance}" for t in result.taxa]) + "Asdda"
        else:
            median_depth = f"Not available for {result.data_source} input data"
        result.notes.append(f"Species detected: {', '.join([s['species'] for s in species_rows])}. Resistance detection is not yet available for this species.")

        context = {
            'd': result.model_dump(),
            'median_depth': median_depth,
            'species_rows': species_rows,
            'notes': '\n\n'.join(result.notes),
            'species': ";".join([s['species'] for s in species_rows]) if len(species_rows)>0 else "Unclassified"
        }

        logging.debug(f"Writing output to {output_filename} using template {self.template_filename}")

        tpl = DocxTemplate(self.template_filename)
        tpl.render(context)
        tpl.save(output_filename)

class DefaultResistanceTemplate(DocxResultTemplate):
    __template_name__ = "default-resistance"
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        __dir__ = os.path.dirname(os.path.abspath(__file__))
        self.template_filename = os.path.join(__dir__, "report-templates", "default_resistance_template.docx")
    def write_output(
            self, 
            result: ProfileResult, 
            conf: dict,
            output_filename: str
        ):
        
        id2name = rv2genes(conf['bed'])

        


        if isinstance(result.qc, VcfQC):
            qc_check = {d:100 for d in id2name.values()}
            result.notes.append("This report was generated from a VCF where gene coverage is not available. It is assumed all genes are covered well sufficiently.")
        else:
            qc_check = {d.target:d.percent_depth_pass for d in result.qc.target_qc}


        gene2drugs = get_gene2drugs(conf['bed'])
        table = [{'name':d, 'genes':[]} for d in conf['drugs']]
        for gene,drugs in gene2drugs.items():
            for d in drugs:
                items = [i for i in table if i['name']==d]
                if len(items)>0:
                    item = items[0]
                    item['genes'].append({'name':gene, 'mechanisms':[]})

        for var in result.dr_variants:
            for d in var.drugs:
                item = [i for i in table if i['name']==d['drug']][0]
                gene_item = [i for i in item['genes'] if i['name']==var.gene_name][0]
                comment = ["Resistance mutation detected"]
                for ann in var.annotation:
                    print(ann)
                    if 'comment' in ann and ann['comment']!="":
                        comment.append(ann['comment'])
                comment = ". ".join(comment)
                gene_item['mechanisms'].append({
                    'mechanism':var.change,
                    'freq':var.freq,
                    'comment':comment
                })

        for dr_gene in result.dr_genes:
            for d in gene2drugs.get(dr_gene.gene_name,[]):
                item = [i for i in table if i['name']==d][0]
                gene_item = [i for i in item['genes'] if i['name']==dr_gene.gene_name][0]
                gene_item['mechanisms'].append({
                    'mechanism':'Resistance gene detected',
                    'freq':'',
                    'comment':'Presence of a functional copy of this gene is associated with resistance to this drug.'
                }) 

        print(table)

        


        rows = []

        for d in table:
            for g in d['genes']:
                if len(g['mechanisms']) == 0:
                    rows.append({
                        'drug':d['name'].replace("_"," ").title(),
                        'gene':g['name'],
                        'qc':qc_check[g['name']],
                        'mechanism':'',
                        'freq':'',
                        'comment':''
                    })
                else:
                    for v in g['mechanisms']:
                        rows.append({
                            'drug':d['name'].replace("_"," ").title(),
                            'gene':g['name'],
                            'qc':qc_check[g['name']],
                            'mechanism':v['mechanism'],
                            'freq':int(v['freq']*100) if v['freq']!='' else '',
                            'comment':v['comment']
                        })
        print(rows)

        resistant_drugs_tmp = [r['drug'] for r in rows if r['mechanism']!=""]
        resistant_drugs = [d for d in conf['drugs'] if d.title() in resistant_drugs_tmp]
        if len(resistant_drugs)==1:
            result_summary = f"Known resistance determinants for {resistant_drugs[0]} detected."
        elif len(resistant_drugs)>1:
            result_summary = f"Known resistance determinants for {', '.join(resistant_drugs[:-1])} and {resistant_drugs[-1]} detected."
        else:
            result_summary = "No known resistance determinants detected."


        poor_coverage_genes = [g for g in qc_check if qc_check[g]<99]
        if len(poor_coverage_genes)>0:
            result.notes.append(f"Insufficient coverage detected in {len(poor_coverage_genes)} genes.")

        subspecies = []
        for d in result.barcode:
            subspecies.append({
                'name':d.id,
                'frequency':d.frequency,
                'info':", ".join(d.info)
            })
        

        species_rows = []
        for t in result.taxa:
            species_rows.append({
                'species': t.species,
                'accession': t.accession,
                'ani': round(t.ani,2) if t.ani else None,
                'rabund': round(t.relative_abundance,2) if t.relative_abundance else None,
                'notes': '\n'.join(t.notes)
            })

        if isinstance(result.qc,FastaQC):
            median_depth = 'Median depth not available for Fasta input data'
        elif isinstance(result.qc,VcfQC):
            median_depth = 'Median depth not available for VCF input data'
        else:
            median_depth = result.qc.target_median_depth


        context = {
            'd': result.model_dump(),
            'result_summary': result_summary,
            'species_rows': species_rows,
            'rows':rows,
            'qc_check': qc_check,
            'notes': '\n\n'.join(result.notes),
            'target_median_depth': median_depth,
            'subspecies': subspecies,
            'has_subspecies': len(subspecies)>0,
            'species': ";".join([s['species'] for s in species_rows]) if len(species_rows)>0 else "Unclassified"
        }

        logging.debug(f"Writing output to {output_filename} using template {self.template_filename}")

        tpl = DocxTemplate(self.template_filename)
        tpl.render(context)
        tpl.save(output_filename)

        merge_cells(output_filename)



def write_docx(result: ProfileResult,conf,outfile,template_file = None, plugin = None):
    # if template_file is None:
    #     if isinstance(result, ProfileResult):
    #         template_file = sys.prefix+"/share/ntm-profiler/default_resistance_template.docx"
    #     else:
    #         template_file = sys.prefix+"/share/ntm-profiler/default_species_template.docx"
    
    if plugin:
        plugin_cls = plugin()
        plugin_cls.write_output(result, conf,outfile)

    else:
        if isinstance(result, ProfileResult):
            output_cls = DefaultResistanceTemplate()
        else:
            output_cls = DefaultSpeciesTemplate()
        output_cls.write_output(result, conf, outfile)
        